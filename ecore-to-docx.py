#!/usr/bin/env python3
"""
ecore_to_docx.py

Convert an Eclipse EMF .ecore file into a Word document containing one table
per EClass. Each table lists the class's structural features (EAttributes and
EReferences) with columns:

    Name | Kind | Type | Cardinality | Default | Description

- "Kind" distinguishes attribute / reference / containment.
- "Type" resolves cross-package references via href (e.g. ecore primitives,
  or other .ecore files in the same directory, if provided).
- "Description" is pulled from EAnnotations: it looks first for a GenModel
  "documentation" detail, then falls back to any other EAnnotation detail
  whose key contains "doc" (case-insensitive), then the first detail value.

Usage:
    python ecore_to_docx.py model.ecore -o data_definitions.docx
    python ecore_to_docx.py *.ecore -o combined.docx
    python ecore_to_docx.py path/to/dir -o combined.docx       # recursive
    python ecore_to_docx.py dir1 file.ecore dir2 -o out.docx   # mixed
    python ecore_to_docx.py model.ecore --title "Data Definition Document"
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# XML namespaces used in .ecore files
# ---------------------------------------------------------------------------
NS = {
    "xmi":   "http://www.omg.org/XMI",
    "xsi":   "http://www.w3.org/2001/XMLSchema-instance",
    "ecore": "http://www.eclipse.org/emf/2002/Ecore",
}

ECORE_PRIMITIVE_RE = re.compile(
    r"http://www\.eclipse\.org/emf/2002/Ecore#//(E[A-Za-z]+)"
)

# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------
@dataclass
class Feature:
    name: str
    kind: str          # "attribute", "reference", or "containment"
    type_str: str
    cardinality: str
    default: str
    description: str


@dataclass
class EClassInfo:
    name: str
    package: str
    is_abstract: bool = False
    is_interface: bool = False
    is_enum: bool = False
    is_datatype: bool = False
    supertypes: list[str] = field(default_factory=list)
    description: str = ""
    features: list[Feature] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------
def _cardinality(lower: str, upper: str) -> str:
    """Format lowerBound/upperBound into a readable cardinality."""
    lo = lower if lower not in (None, "") else "0"
    up = upper if upper not in (None, "") else "1"
    if up == "-1":
        up = "*"
    if lo == up:
        return lo
    return f"{lo}..{up}"


def _resolve_type(etype_attr: Optional[str], href_attr: Optional[str]) -> str:
    """
    Turn an eType reference into a human-readable type name.
    Handles three forms:
      1. eType="#//SomeClass"           -> SomeClass (same file)
      2. <eType href="...#//EString"/>  -> EString   (ecore primitives or cross-file)
      3. eType="ecore:EDataType ...#//EString"
    """
    ref = etype_attr or href_attr or ""
    if not ref:
        return ""

    # Ecore primitive via full URL
    m = ECORE_PRIMITIVE_RE.search(ref)
    if m:
        return m.group(1)

    # Fragment-style: "#//Something" or "some.ecore#//Something"
    if "#//" in ref:
        frag = ref.split("#//", 1)[1]
        # Could be nested like "Package/Class" — take the last segment
        return frag.split("/")[-1]

    # Fallback: strip any URI prefix
    return ref.rsplit("/", 1)[-1]


def _extract_description(element: ET.Element) -> str:
    """Look for documentation-style EAnnotations on an element."""
    for ann in element.findall("eAnnotations"):
        details = ann.findall("details")
        # Priority 1: a detail with key exactly "documentation"
        for d in details:
            if d.get("key", "").lower() == "documentation":
                return (d.get("value") or "").strip()
        # Priority 2: any detail whose key contains "doc"
        for d in details:
            if "doc" in d.get("key", "").lower():
                return (d.get("value") or "").strip()
        # Priority 3: first detail's value
        if details:
            return (details[0].get("value") or "").strip()
    return ""


def _parse_feature(elem: ET.Element) -> Feature:
    """Parse an eStructuralFeatures element (attribute or reference)."""
    xsi_type = elem.get(f"{{{NS['xsi']}}}type", "")
    is_ref = xsi_type.endswith("EReference")
    containment = elem.get("containment", "false").lower() == "true"

    if is_ref:
        kind = "containment" if containment else "reference"
    else:
        kind = "attribute"

    # Type can be inline attribute OR a nested <eType href="..."/>
    etype_attr = elem.get("eType")
    href_attr = None
    etype_child = elem.find("eType")
    if etype_child is not None:
        href_attr = etype_child.get(f"{{{NS['xmi']}}}href") or etype_child.get("href")

    type_str = _resolve_type(etype_attr, href_attr)

    cardinality = _cardinality(elem.get("lowerBound"), elem.get("upperBound"))
    default = elem.get("defaultValueLiteral", "") or ""
    description = _extract_description(elem)

    return Feature(
        name=elem.get("name", ""),
        kind=kind,
        type_str=type_str,
        cardinality=cardinality,
        default=default,
        description=description,
    )


def _parse_supertypes(eclass: ET.Element) -> list[str]:
    raw = eclass.get("eSuperTypes", "")
    if not raw:
        return []
    names = []
    for token in raw.split():
        names.append(_resolve_type(token, None))
    return names


def parse_ecore(path: Path) -> list[EClassInfo]:
    """Parse a .ecore file and return all EClasses within all EPackages."""
    tree = ET.parse(path)
    root = tree.getroot()

    # Root may itself be an EPackage, or an <xmi:XMI> wrapping multiple packages
    if root.tag.endswith("EPackage"):
        packages = [root]
    else:
        packages = root.findall(".//{%s}EPackage" % NS["ecore"]) or root.findall(".//EPackage")
        if not packages:
            packages = [root]

    classes: list[EClassInfo] = []
    for pkg in packages:
        pkg_name = pkg.get("name", path.stem)
        for classifier in pkg.findall("eClassifiers"):
            xsi_type = classifier.get(f"{{{NS['xsi']}}}type", "")

            if xsi_type.endswith("EClass"):
                info = EClassInfo(
                    name=classifier.get("name", ""),
                    package=pkg_name,
                    is_abstract=classifier.get("abstract", "false").lower() == "true",
                    is_interface=classifier.get("interface", "false").lower() == "true",
                    supertypes=_parse_supertypes(classifier),
                    description=_extract_description(classifier),
                )
                for feat_elem in classifier.findall("eStructuralFeatures"):
                    info.features.append(_parse_feature(feat_elem))
                classes.append(info)

            elif xsi_type.endswith("EEnum"):
                info = EClassInfo(
                    name=classifier.get("name", ""),
                    package=pkg_name,
                    is_enum=True,
                    description=_extract_description(classifier),
                )
                # Each literal becomes a row. "Default" column holds the
                # integer value; "Description" uses GenModel docs if present.
                for lit in classifier.findall("eLiterals"):
                    lit_name = lit.get("name", "")
                    literal_str = lit.get("literal") or lit_name
                    value = lit.get("value", "")
                    info.features.append(Feature(
                        name=literal_str,
                        kind="literal",
                        type_str="",
                        cardinality="",
                        default=value,
                        description=_extract_description(lit),
                    ))
                classes.append(info)

            elif xsi_type.endswith("EDataType"):
                # Custom data types (usually wrapping a Java class). They
                # have no features or literals, so we render them as a
                # single-row "data type" block that documents what they
                # alias to.
                dt_name = classifier.get("name", "")
                instance_class = (
                    classifier.get("instanceClassName")
                    or classifier.get("instanceTypeName")
                    or ""
                )
                serializable = classifier.get("serializable", "true").lower() == "true"
                info = EClassInfo(
                    name=dt_name,
                    package=pkg_name,
                    is_datatype=True,
                    description=_extract_description(classifier),
                )
                info.features.append(Feature(
                    name=dt_name,
                    kind="datatype",
                    type_str=instance_class,
                    cardinality="",
                    default="" if serializable else "not serializable",
                    description=_extract_description(classifier),
                ))
                classes.append(info)

            # Other classifier kinds (if any) are still skipped.

    return classes


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------
HEADER_FILL = "2E75B6"   # blue header
ROW_ALT_FILL = "F2F2F2"  # light grey zebra
BORDER_COLOR = "BFBFBF"

COLUMNS = [
    ("Class",       Cm(3.0)),
    ("Name",        Cm(3.0)),
    ("Kind",        Cm(2.2)),
    ("Type",        Cm(3.0)),
    ("Default",     Cm(2.0)),
    ("Description", Cm(6.0)),
]


def _shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), BORDER_COLOR)
        borders.append(b)
    tc_pr.append(borders)


def _write_cell(cell, text: str, *, bold: bool = False, italic: bool = False,
                color: Optional[RGBColor] = None, size: int = 10) -> None:
    cell.text = ""  # clear default paragraph
    para = cell.paragraphs[0]
    run = para.add_run(text or "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_borders(cell)


def _add_class_heading(doc: Document, cls: EClassInfo) -> None:
    heading = doc.add_heading(level=3)
    run = heading.add_run(cls.name)
    run.font.size = Pt(12)

    # Subtitle line with package + modifiers + supertypes
    bits = [f"Package: {cls.package}"]
    if cls.is_abstract:
        bits.append("abstract")
    if cls.is_interface:
        bits.append("interface")
    if cls.supertypes:
        bits.append("extends " + ", ".join(cls.supertypes))
    sub = doc.add_paragraph()
    sub_run = sub.add_run(" · ".join(bits))
    sub_run.italic = True
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    if cls.description:
        desc = doc.add_paragraph(cls.description)
        desc.paragraph_format.space_after = Pt(4)


def _add_file_table(doc: Document, classes: list[EClassInfo]) -> None:
    """
    Render one combined table for a single .ecore file.
    Columns: Class | Name | Kind | Type | Default | Description.
    Each class contributes its features as rows; the Class cell shows the
    class name only on the first row of that class's block for readability.
    Classes with no features still get a single row so they're visible.
    """
    # Count total data rows so we can size the table up-front
    total_rows = 0
    for cls in sorted(classes, key=lambda c: c.name.lower()):
        total_rows += max(len(cls.features), 1)

    if total_rows == 0:
        note = doc.add_paragraph()
        r = note.add_run("(No EClasses found in this file.)")
        r.italic = True
        r.font.size = Pt(9)
        return

    table = doc.add_table(rows=1 + total_rows, cols=len(COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    for i, (_, width) in enumerate(COLUMNS):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    header = table.rows[0]
    for i, (label, _) in enumerate(COLUMNS):
        cell = header.cells[i]
        _write_cell(cell, label, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        _shade_cell(cell, HEADER_FILL)

    # Data rows
    r_idx = 1
    class_block = 0  # toggles zebra shading per class block, not per row
    for cls in sorted(classes, key=lambda c: c.name.lower()):
        feats = cls.features if cls.features else [None]
        shade = (class_block % 2 == 1)
        for f_pos, feat in enumerate(feats):
            row = table.rows[r_idx]
            class_label = cls.name if f_pos == 0 else ""
            if feat is None:
                values = [class_label, "", "", "", "", "(no features)"]
            else:
                values = [
                    class_label,
                    feat.name,
                    feat.kind,
                    feat.type_str,
                    feat.default,
                    feat.description,
                ]
            for c_idx, val in enumerate(values):
                cell = row.cells[c_idx]
                is_class_col = (c_idx == 0)
                _write_cell(
                    cell, val,
                    bold=(is_class_col and f_pos == 0),
                    italic=(feat is None and c_idx == len(values) - 1),
                    size=9,
                )
                if shade:
                    _shade_cell(cell, ROW_ALT_FILL)
            r_idx += 1
        class_block += 1

    doc.add_paragraph()  # spacing after the table


def build_document(
    classes_by_file: dict[str, list[EClassInfo]],
    title: str,
) -> Document:
    doc = Document()

    # Landscape helps the wider tables breathe
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Title
    title_p = doc.add_heading(title, level=0)
    for run in title_p.runs:
        run.font.size = Pt(22)

    intro = doc.add_paragraph(
        "This document defines the structural classes extracted from the provided "
        "Ecore model(s). Each table below corresponds to one .ecore source file "
        "and lists every structural feature (attribute or reference) of every "
        "EClass in that file."
    )
    intro.paragraph_format.space_after = Pt(12)

    for source_file, classes in classes_by_file.items():
        h1 = doc.add_heading(level=1)
        h1.add_run(Path(source_file).name)

        # Subtitle: the packages contained in this file
        packages = sorted({c.package for c in classes})
        if packages:
            sub = doc.add_paragraph()
            sr = sub.add_run("Package(s): " + ", ".join(packages))
            sr.italic = True
            sr.font.size = Pt(9)
            sr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

        if not classes:
            note = doc.add_paragraph()
            r = note.add_run("(No EClasses in this file.)")
            r.italic = True
            r.font.size = Pt(9)
            continue

        _add_file_table(doc, classes)

    return doc


# ---------------------------------------------------------------------------
# Input expansion
# ---------------------------------------------------------------------------
def _walk_for_ecore(root: Path) -> list[Path]:
    """
    Recursively find *.ecore files under `root`, skipping subdirectories we
    don't have permission to read instead of aborting the whole walk.
    `os.walk` with onerror lets us log and continue; `rglob` would raise.
    """
    found: list[Path] = []

    def _on_error(err: OSError) -> None:
        # err.filename is the path that couldn't be read
        print(f"  WARNING: skipping unreadable path: {err.filename} "
              f"({err.strerror or err})", file=sys.stderr)

    for dirpath, dirnames, filenames in os.walk(root, onerror=_on_error,
                                                followlinks=False):
        # Proactively skip common Windows protected/system dirs that we
        # definitely don't want to descend into (and which would otherwise
        # trigger PermissionError and noise up the output).
        dirnames[:] = [
            d for d in dirnames
            if d.lower() not in {
                "system volume information",
                "$recycle.bin",
                ".git",
                "node_modules",
            }
        ]
        for fn in filenames:
            if fn.lower().endswith(".ecore"):
                found.append(Path(dirpath) / fn)

    return sorted(found)


def collect_ecore_files(inputs: list[str]) -> list[Path]:
    """
    Expand CLI inputs into a deduplicated, sorted list of .ecore file paths.
    - A file path is included as-is (must end in .ecore).
    - A directory is walked recursively for *.ecore files, skipping any
      subdirectory we lack permission to read.
    """
    seen: set[Path] = set()
    result: list[Path] = []

    for raw in inputs:
        p = Path(raw)
        if not p.exists():
            print(f"ERROR: path not found: {p}", file=sys.stderr)
            sys.exit(2)

        if p.is_dir():
            try:
                found = _walk_for_ecore(p)
            except PermissionError as e:
                # The top-level directory itself is unreadable.
                print(f"ERROR: cannot read directory {p}: {e.strerror or e}",
                      file=sys.stderr)
                continue
            if not found:
                print(f"  (no .ecore files under {p})")
            for f in found:
                try:
                    resolved = f.resolve()
                except (OSError, PermissionError):
                    resolved = f.absolute()
                if resolved not in seen:
                    seen.add(resolved)
                    result.append(f)
        elif p.is_file():
            if p.suffix.lower() != ".ecore":
                print(f"WARNING: skipping non-.ecore file: {p}", file=sys.stderr)
                continue
            try:
                resolved = p.resolve()
            except (OSError, PermissionError):
                resolved = p.absolute()
            if resolved not in seen:
                seen.add(resolved)
                result.append(p)

    return result


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main() -> int:
    ap = argparse.ArgumentParser(
        description="Convert .ecore file(s) into a Word document of data-definition tables. "
                    "Accepts individual files and/or directories (searched recursively)."
    )
    ap.add_argument("inputs", nargs="+",
                    help="One or more .ecore files and/or directories to search recursively")
    ap.add_argument("-o", "--output", default="ecore_tables.docx",
                    help="Output .docx path (default: ecore_tables.docx)")
    ap.add_argument("--title", default="Data Definitions",
                    help="Document title (default: 'Data Definitions')")
    args = ap.parse_args()

    files = collect_ecore_files(args.inputs)
    if not files:
        print("No .ecore files found.", file=sys.stderr)
        return 1

    print(f"Found {len(files)} .ecore file(s):")
    classes_by_file: dict[str, list[EClassInfo]] = {}
    total = 0
    for p in files:
        try:
            classes = parse_ecore(p)
        except ET.ParseError as e:
            print(f"  ERROR parsing {p}: {e}", file=sys.stderr)
            continue
        except PermissionError as e:
            print(f"  ERROR: permission denied reading {p}: {e.strerror or e}",
                  file=sys.stderr)
            continue
        except OSError as e:
            print(f"  ERROR reading {p}: {e.strerror or e}", file=sys.stderr)
            continue
        classes_by_file[str(p)] = classes
        total += len(classes)
        print(f"  {p}: {len(classes)} EClass(es)")

    if total == 0:
        print("No EClasses found in the provided file(s).", file=sys.stderr)
        return 1

    doc = build_document(classes_by_file, title=args.title)
    out_path = Path(args.output)
    try:
        doc.save(out_path)
    except PermissionError as e:
        # Most common cause on Windows: the output .docx is open in Word,
        # or the containing folder is read-only / syncing.
        from datetime import datetime
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        fallback = out_path.with_name(f"{out_path.stem}_{stamp}{out_path.suffix}")
        print(
            f"\nERROR: could not write to {out_path}\n"
            f"  ({e.strerror or e})\n"
            f"  This usually means the file is open in Word, or the folder is\n"
            f"  read-only / protected. Close the document and retry, or pick a\n"
            f"  different -o path.\n"
            f"  Attempting fallback: {fallback}",
            file=sys.stderr,
        )
        try:
            doc.save(fallback)
            print(f"\nWrote {total} table(s) to {fallback}")
            return 0
        except PermissionError as e2:
            print(f"  Fallback also failed: {e2.strerror or e2}", file=sys.stderr)
            return 2
    print(f"\nWrote {total} table(s) to {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
